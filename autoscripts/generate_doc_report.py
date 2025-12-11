import os
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Please install it with 'pip install python-docx'")
    exit(1)

def main():
    document = Document()

    # Title
    head = document.add_heading('Geospatial Tagging Pipeline & Methodology', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('Automated Pipeline Report for Multi-City Expansion')
    
    # Section 1: Pipeline Overview
    document.add_heading('1. Pipeline Overview', level=1)
    
    intro = document.add_paragraph(
        "The pipeline is designed to collect, process, and classify Volunteered Geographic Information (VGI) "
        "from OpenStreetMap (OSM) for multiple cities in Pakistan (Islamabad, Lahore, Karachi). "
        "The goal is to deduce place types from unstructured text descriptions using Natural Language Processing (NLP) "
        "and Machine Learning models."
    )

    # 1.1 Data Collection
    document.add_heading('1.1 Data Collection', level=2)
    p = document.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('01_collect_data_multicity.py')
    document.add_paragraph(
        "This step queries the Overpass API to fetch OSM nodes, ways, and relations within configured bounding boxes. "
        "It iterates through a list of cities defined in 'cities.yaml'. "
        "Raw JSON responses are saved and then flattened into a CSV format containing text descriptions (name, tags) and metadata."
    )

    # 1.2 Preprocessing
    document.add_heading('1.2 Preprocessing', level=2)
    p = document.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('02_preprocess_data_multicity.py')
    document.add_paragraph(
        "This step merges data from all cities and performs text cleaning. "
        "It handles Urdu characters, unescapes HTML entities, and synthesizes a final 'description' "
        "by combining the place name and its tags if a raw description is missing."
    )

    # 1.3 Pattern Mining
    document.add_heading('1.3 Pattern Mining', level=2)
    p = document.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('03_pattern_mining.py')
    document.add_paragraph(
        "Frequent Pattern Mining (similar to Apriori/FP-Growth) is applied to tokenized descriptions "
        "to identify common terms and co-occurring words. This helps in understanding the vocabulary structure "
        "and validating potential feature keywords."
    )

    # 1.4 Embedding & Classification (Baseline)
    document.add_heading('1.4 Embedding & Classification', level=2)
    p = document.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('04_embedding_classification.py')
    document.add_paragraph(
        "Text descriptions are converted into dense vector representations using Sentence-BERT (SBERT). "
        "A Logistic Regression model is trained as a baseline to benchmark classification performance."
    )

    # 1.5 Bayesian Elastic Net
    document.add_heading('1.5 Bayesian Elastic Net', level=2)
    p = document.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('05_bayesian_net.py')
    document.add_paragraph(
        "The core model of this thesis. It uses an Elastic Net penalty (L1 + L2) to induce sparsity "
        "while handling correlated features. It is implemented via SGDClassifier with log-loss, "
        "approximating a Bayesian approach where the regularization terms correspond to prior distributions."
    )

    # 1.6 Explainable AI (XAI)
    document.add_heading('1.6 Explainability', level=2)
    p = document.add_paragraph()
    p.add_run('Script: ').bold = True
    p.add_run('06_xai.py')
    document.add_paragraph(
        "SHAP (SHapley Additive exPlanations) values are calculated to interpret the model's predictions. "
        "This reveals which words or semantic features contribute most to specific classifications."
    )

    # Section 2: Mathematical Methodology
    document.add_heading('2. Mathematical Methodology: Bayesian Elastic Net', level=1)

    document.add_paragraph(
        "The 'Bayesian Elastic Net' in this context refers to a regularized Logistic Regression model "
        "that can be interpreted probabilistically."
    )

    # 2.1 Objective Function
    document.add_heading('2.1 Objective Function', level=2)
    document.add_paragraph(
        "We minimize the regularized negative log-likelihood (Log Loss) with Elastic Net regularization. "
        "For a dataset of N samples, the objective function L(w) is:"
    )
    
    # Simple math representation in text
    document.add_paragraph(
        "L(w) = -(1/N) * Σ [y_i * log(p_i) + (1 - y_i) * log(1 - p_i)] + λ [ α ||w||_1 + (1 - α)/2 ||w||_2^2 ]"
    ).style = 'Quote'

    document.add_paragraph(
        "Where:\n"
        "- w consists of the model weights/coefficients.\n"
        "- p_i is the predicted probability for sample i (via Sigmoid function).\n"
        "- λ (alpha in code) controls the total regularization strength.\n"
        "- α (l1_ratio in code, implies mixing) controls the balance between L1 (Lasso) and L2 (Ridge)."
    )

    # 2.2 Bayesian Interpretation
    document.add_heading('2.2 Bayesian Interpretation', level=2)
    document.add_paragraph(
        "Regularization terms can be theoretically derived from Bayesian priors on the weights:"
    )
    document.add_paragraph(
        "1. L2 Regularization (Ridge) corresponds to placing a Gaussian Prior on the weights.\n"
        "   P(w) ~ Normal(0, σ²)"
    )
    document.add_paragraph(
        "2. L1 Regularization (Lasso) corresponds to placing a Laplace Prior on the weights.\n"
        "   P(w) ~ Laplace(0, b)"
    )
    document.add_paragraph(
        "The Elastic Net combines both, effectively assuming a prior that encourages both sparsity (Laplace) "
        "and grouping/smoothness (Gaussian). Minimizing the objective function is equivalent to finding the "
        "Maximum A Posteriori (MAP) estimate of the weights."
    )

    # 2.3 Optimization
    document.add_heading('2.3 SGD Optimization', level=2)
    document.add_paragraph(
        "Due to the scale of data and high dimensionality of embedding vectors, we use Stochastic Gradient Descent (SGD). "
        "SGD updates weights iteratively using small batches of data:"
    )
    document.add_paragraph(
        "w_{t+1} = w_t - η * (∇Loss(w_t) + ∇Regularization(w_t))"
    ).style = 'Quote'
    document.add_paragraph(
        "This allows for efficient training on large multi-city datasets without requiring the entire dataset to fit into memory for matrix inversion."
    )

    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 
                               "Geospatial_Tagging_Pipeline_Report.docx")
    document.save(output_path)
    print(f"Report generated successfully at: {output_path}")

if __name__ == "__main__":
    main()
