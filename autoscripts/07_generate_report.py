import os
import json
import logging
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Configuration ---
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATAREPORTED_DIR = os.path.join(BASE_DIR, "datareported")
CLASSIFICATION_DIR = os.path.join(DATAREPORTED_DIR, "classification")
XAI_DIR = os.path.join(DATAREPORTED_DIR, "xai")
BEN_FEATURES_FILE = os.path.join(DATAREPORTED_DIR, "ben_features", "ben_selected_features.json")
RESULTS_FILE = os.path.join(CLASSIFICATION_DIR, "benchmark_results.json")
OUTPUT_REPORT = os.path.join(BASE_DIR, "Extended_Pipeline_Report.docx")

# Logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(message)s")

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_para(doc, text):
    doc.add_paragraph(text)

def main():
    logging.info("Generating Report (Autoscript 07)...")
    
    doc = Document()
    doc.add_heading('Extended Geospatial Tagging Pipeline Report', 0)
    
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Load Stats to report
    raw_files = [f for f in os.listdir(os.path.join(DATAREPORTED_DIR, "raw")) if f.endswith(".csv")]
    
    # 1. Methodology
    add_heading(doc, "1. Methodology", level=1)
    
    add_heading(doc, "1.1 Data Collection & Preprocessing", level=2)
    add_para(doc, 
        f"Data was collected for {len(raw_files)} cities (Islamabad, Lahore, Karachi, Rawalpindi, Kohat, Peshawar, Attock, Chitral, Skardu, Quetta, Sialkot, Muzaffarabad) using the Overpass API. "
        "The preprocessing pipeline included Roman Urdu normalization and POS tagging."
    )
    
    # Preprocessing Stats
    if os.path.exists(os.path.join(DATAREPORTED_DIR, "preprocessed", "merged_dataset.csv")):
         merged_df = pd.read_csv(os.path.join(DATAREPORTED_DIR, "preprocessed", "merged_dataset.csv"))
         add_para(doc, f"Total Preprocessed Records: {len(merged_df)}")
    
    add_heading(doc, "1.2 Pattern Mining (Guided FP-Growth)", level=2)
    patterns_file = os.path.join(DATAREPORTED_DIR, "fp_growth", "guided_patterns.csv")
    num_patterns = len(pd.read_csv(patterns_file)) if os.path.exists(patterns_file) else "N/A"
    
    add_para(doc,
        f"We employed Guided FP-Growth to mine frequent itemsets. A total of {num_patterns} relevant patterns were identified."
    )
    
    add_heading(doc, "1.3 Feature Selection (Bayesian Elastic Net)", level=2)
    ben_feats_file = os.path.join(DATAREPORTED_DIR, "ben_features", "ben_selected_features.json")
    num_feats = 0
    if os.path.exists(ben_feats_file):
        with open(ben_feats_file, 'r') as f:
            num_feats = len(json.load(f).get('features', []))
            
    add_para(doc,
        f"Bayesian Elastic Net (BEN) regularization selected {num_feats} high-utility features from the mined patterns."
    )
    
    add_heading(doc, "2. Mathematical Methodology: Bayesian Elastic Net", level=1)
    # Mathematical formulation (Aligned with Request)
    add_heading(doc, '2.1 Objective Function', level=2)
    
    add_para(doc, "The objective function minimized is (Negative Log-Likelihood + Elastic Net Prior):")
    add_para(doc, "L(w) = -(1/N) * Σ [y_i * log(p_i) + (1 - y_i) * log(1 - p_i)] + λ [ α ||w||_1 + (1 - α)/2 ||w||_2^2 ]")
    
    add_para(doc,
        "Where:\n"
        "- p_i is the predicted probability.\n"
        "- λ (alpha in code) controls regularization strength.\n"
        "- α (l1_ratio) balances L1 (Lasso) and L2 (Ridge) penalties."
    )
    
    add_heading(doc, "2.2 Bayesian Interpretation", level=2)
    add_para(doc,
        "1. L2 (Ridge): Corresponds to a Gaussian Prior P(w) ~ Normal(0, σ²).\n"
        "2. L1 (Lasso): Corresponds to a Laplace Prior P(w) ~ Laplace(0, b).\n"
        "The Elastic Net combines both, encouraging sparsity and grouping."
    )

    add_heading(doc, "3. Hybrid Classification", level=1)
    add_para(doc,
        "The final classification model is a Hybrid Architecture fusing:"
    )
    doc.add_paragraph("- Semantic Embeddings: DistilBERT (Contextual understanding)", style='List Bullet')
    doc.add_paragraph("- Linguistic Features: Selected via BEN (Domain-specific signals)", style='List Bullet')
    add_para(doc, "These inputs are concatenated and passed through a Multi-Layer Perceptron (MLP) for final tag prediction.")

    # 2. Results
    add_heading(doc, "2. Results & Benchmarking", level=1)
    
    if os.path.exists(RESULTS_FILE):
        with open(RESULTS_FILE, "r") as f:
            res = json.load(f)
            
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        metrics = [
            ("Baseline Accuracy (Word2Vec)", f"{res['baseline_word2vec_accuracy']:.4f}"),
            ("Hybrid Model Accuracy (DistilBERT+BEN)", f"{res['hybrid_model_accuracy']:.4f}"),
            ("Improvement", f"+{res['improvement']:.4f}"),
            ("Total Classes", str(res['classes']))
        ]
        
        for k, v in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = k
            row_cells[1].text = v
            
        add_para(doc, "\n")
        add_para(doc, f"The Hybrid model achieved an accuracy of {res['hybrid_model_accuracy']:.2%}, outperforming the baseline by {res['improvement']:.2%}.")
    else:
        add_para(doc, "Results file not found.")

    # 3. Explainability
    add_heading(doc, "3. Explainable AI (SHAP)", level=1)
    
    shap_plot = os.path.join(XAI_DIR, "shap_summary_plot.png")
    if os.path.exists(shap_plot):
        doc.add_picture(shap_plot, width=Inches(6))
        add_para(doc, "Figure 1: SHAP Summary Plot showing top features influencing the model predictions.")
    else:
        add_para(doc, "SHAP plot not found.")
        
    doc.save(OUTPUT_REPORT)
    logging.info(f"Report saved to {OUTPUT_REPORT}")

if __name__ == "__main__":
    main()
